# app/components/recruitment_admin.py

from __future__ import annotations

import os
import base64
import json
import uuid
from pathlib import Path
from typing import Any, Dict, List

import requests
import streamlit as st
from docx import Document

API_BASE_URL = os.getenv("API_BASE_URL", "http://localhost:9898/api/v1")


def _get(url: str, timeout: int = 30) -> requests.Response:
    return requests.get(url, timeout=timeout)


def _post_multipart(url: str, data: Dict[str, Any], files: Dict[str, Any], timeout: int = 120) -> requests.Response:
    return requests.post(url, data=data, files=files, timeout=timeout)


def _patch(url: str, data: Dict[str, Any], timeout: int = 30) -> requests.Response:
    return requests.patch(url, data=data, timeout=timeout)


def fetch_admin_recruitments() -> List[Dict[str, Any]]:
    resp = _get(f"{API_BASE_URL}/recruitments/admin/list", timeout=60)
    if resp.status_code != 200:
        st.error(f"채용공고 목록 조회 실패: {resp.status_code}")
        return []
    return resp.json()


def upload_recruitment(data: Dict[str, Any], file) -> Dict[str, Any] | None:
    # None 값 제거 및 문자열 필드 정리
    clean_data = {}
    for key, value in data.items():
        if value is not None:
            if isinstance(value, str):
                clean_data[key] = value.strip()
            else:
                clean_data[key] = value
    
    files = {"file": (file.name, file.getvalue(), file.type)}
    resp = _post_multipart(f"{API_BASE_URL}/recruitments/admin/upload", data=clean_data, files=files, timeout=180)
    if resp.status_code != 200:
        try:
            error_detail = resp.json()
            error_msg = error_detail.get("detail", resp.text)
        except Exception:
            error_msg = resp.text
        st.error(f"등록 실패 ({resp.status_code}): {error_msg}")
        return None
    return resp.json()


def render_recruit_admin_page() -> None:
    st.title("📂 채용공고 관리")
    
    # 등록 성공 후 상단으로 스크롤 (한 번만 실행)
    if st.session_state.get("recruit_admin_scroll_to_top", False):
        st.markdown(
            """
            <script>
            window.scrollTo({ top: 0, behavior: 'smooth' });
            </script>
            """,
            unsafe_allow_html=True,
        )
        st.session_state["recruit_admin_scroll_to_top"] = False
    
    # 파일 보기 닫기 후 해당 채용공고 카드로 스크롤 이동
    scroll_to_rec_id = st.session_state.get("recruit_admin_scroll_to_rec_id")
    if scroll_to_rec_id:
        st.markdown(
            f"""
            <script>
            (function() {{
                function scrollToCard() {{
                    const card = document.getElementById('rec-card-{scroll_to_rec_id}');
                    if (card) {{
                        card.scrollIntoView({{ behavior: 'smooth', block: 'center' }});
                        // 하이라이트 효과
                        card.style.transition = 'box-shadow 0.3s ease';
                        card.style.boxShadow = '0 0 0 3px rgba(59, 130, 246, 0.5)';
                        setTimeout(function() {{
                            card.style.boxShadow = '';
                        }}, 2000);
                        return true;
                    }}
                    return false;
                }}
                // 즉시 시도
                if (!scrollToCard()) {{
                    // DOM이 준비되지 않았으면 재시도
                    setTimeout(scrollToCard, 200);
                    setTimeout(scrollToCard, 500);
                    setTimeout(scrollToCard, 1000);
                }}
            }})();
            </script>
            """,
            unsafe_allow_html=True,
        )
        # 플래그 제거
        st.session_state["recruit_admin_scroll_to_rec_id"] = None

    # 목록 헤더 + 등록 버튼
    col_head, col_btn = st.columns([3, 1])
    with col_head:
        st.markdown("#### 📋 채용공고 목록")
    with col_btn:
        if st.button("➕ 채용공고 등록", use_container_width=True, key="toggle_upload", type="primary"):
            st.session_state["recruit_admin_show_upload"] = True
            st.session_state["recruit_admin_list_expanded"] = False  # 목록 접기
            st.rerun()

    # 목록 펼침/접힘 상태 관리 (기본값: True)
    list_expanded = st.session_state.get("recruit_admin_list_expanded", True)
    
    # 새로고침을 위해 캐시 무효화 (등록 후)
    last_upload = st.session_state.get("recruit_admin_last_upload")
    recs = fetch_admin_recruitments()
    
    # 채용공고 목록을 expander로 감싸기
    with st.expander("", expanded=list_expanded):
        if not recs:
            st.info("등록된 채용공고가 없습니다.")
        else:
            # 카드 스타일 정의
            st.markdown(
            """
            <style>
            .rec-card-container {
                padding: 14px 18px;
                border-radius: 12px;
                border: 1px solid rgba(148, 163, 184, 0.2);
                background: linear-gradient(135deg, rgba(15, 23, 42, 0.95) 0%, rgba(30, 41, 59, 0.95) 100%);
                margin-bottom: 6px;
                box-shadow: 0 2px 4px -1px rgba(0, 0, 0, 0.1), 0 1px 2px -1px rgba(0, 0, 0, 0.06);
            }
            .rec-title {
                font-size: 1.2rem;
                font-weight: 700;
                color: #f1f5f9;
                margin-top: 6px;
                margin-bottom: 8px;
                line-height: 1.3;
            }
            .rec-meta {
                font-size: 0.85rem;
                color: #cbd5e1;
                margin: 3px 0;
                line-height: 1.5;
            }
            .rec-meta-label {
                color: #94a3b8;
                font-weight: 500;
            }
            .status-badge {
                display: inline-block;
                padding: 4px 10px;
                border-radius: 8px;
                font-weight: 600;
                font-size: 0.75rem;
                text-align: center;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
                margin-left: 12px;
                vertical-align: middle;
            }
            .action-section {
                background: linear-gradient(135deg, rgba(30, 41, 59, 0.6) 0%, rgba(15, 23, 42, 0.6) 100%);
                border-radius: 10px;
                padding: 10px 14px 10px 14px;
                border: 1px solid rgba(148, 163, 184, 0.2);
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            .action-label {
                font-size: 0.8rem;
                color: #cbd5e1;
                font-weight: 600;
                margin-bottom: 8px;
                margin-top: 6px;
                letter-spacing: 0.3px;
            }
            </style>
            """,
            unsafe_allow_html=True,
            )
            
            status_colors = {"OPEN": "#10b981", "CLOSED": "#f59e0b", "ARCHIVED": "#94a3b8"}
            
            for rec in recs:
                # 카드 컨테이너 (고유 ID 부여)
                rec_id = rec.get("id", "")
                st.markdown(f'<div id="rec-card-{rec_id}" class="rec-card-container">', unsafe_allow_html=True)
                
                col_left, col_right = st.columns([2.5, 1.5])
                
                with col_left:
                    # 제목 + 상태 배지 (같은 라인)
                    status = rec.get("status", "OPEN")
                    color = status_colors.get(status, "#94a3b8")
                    title_with_status = f'<div class="rec-title">{rec.get("title", "-")}<span class="status-badge" style="background: {color}; color: white;">{status}</span></div>'
                    st.markdown(title_with_status, unsafe_allow_html=True)
                    
                    # 메타 정보
                    meta_parts = []
                    if rec.get("job_family"):
                        meta_parts.append(rec["job_family"])
                    if rec.get("role_category"):
                        meta_parts.append(rec["role_category"])
                    if rec.get("location"):
                        meta_parts.append(rec["location"])
                    
                    if meta_parts:
                        meta_display = " | ".join(meta_parts)
                        st.markdown(
                            f'<div class="rec-meta">{meta_display}</div>',
                            unsafe_allow_html=True,
                        )
                    
                    # 기간
                    st.markdown(
                        f'<div class="rec-meta"><span class="rec-meta-label">기간:</span> {rec.get("start_date", "-")} ~ {rec.get("end_date", "-")}</div>',
                        unsafe_allow_html=True,
                    )
                    
                    # 지원자 수 (최근 지원 포함)
                    applicant_count = rec.get("applicant_count", 0)
                    last_application = rec.get("last_application_at", "")
                    if last_application:
                        applicant_text = f'<div class="rec-meta"><span class="rec-meta-label">지원자 수:</span> {applicant_count} <span style="color: #94a3b8;">(최근 지원: {last_application})</span></div>'
                    else:
                        applicant_text = f'<div class="rec-meta"><span class="rec-meta-label">지원자 수:</span> {applicant_count}</div>'
                    st.markdown(applicant_text, unsafe_allow_html=True)
                
                with col_right:
                    # 상태 변경 섹션
                    st.markdown('<div class="action-label">상태 변경</div>', unsafe_allow_html=True)
                    col_select, col_apply = st.columns([2, 1])
                    
                    with col_select:
                        new_status = st.selectbox(
                            "상태 선택",
                            options=["OPEN", "CLOSED", "ARCHIVED"],
                            index=["OPEN", "CLOSED", "ARCHIVED"].index(status),
                            key=f"status_sel_{rec['id']}",
                            label_visibility="collapsed",
                        )
                    
                    with col_apply:
                        if st.button("✅ 적용", key=f"apply_status_{rec['id']}", use_container_width=True, type="secondary"):
                            resp = _patch(
                                f"{API_BASE_URL}/recruitments/admin/{rec['id']}/status",
                                {"status": new_status},
                                timeout=30,
                            )
                            if resp.status_code != 200:
                                st.error(f"상태 변경 실패: {resp.status_code}")
                            else:
                                st.success("상태가 변경되었습니다.")
                                st.rerun()
                    
                    # 채용공고보기 버튼
                    if st.button("📄 채용공고보기", key=f"view_{rec['id']}", use_container_width=True):
                        st.session_state["recruit_admin_view"] = rec
                        st.session_state["recruit_admin_list_expanded"] = False  # 목록 접기
                        st.session_state["recruit_admin_scroll_to_viewer"] = True
                        st.rerun()
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)

    # 선택한 파일 미리보기
    rec_view = st.session_state.get("recruit_admin_view")
    if rec_view:
        st.markdown("---")
        
        # 파일 뷰어 앵커 (항상 배치)
        st.markdown('<div id="file-viewer-anchor"></div>', unsafe_allow_html=True)
        
        # 파일 뷰어로 스크롤 이동 (한 번만 실행)
        scroll_to_viewer = st.session_state.get("recruit_admin_scroll_to_viewer", False)
        if scroll_to_viewer:
            st.markdown(
                """
                <script>
                (function() {
                    function scrollToViewer() {
                        const anchor = document.getElementById('file-viewer-anchor');
                        if (anchor) {
                            // 요소의 위치 계산
                            const rect = anchor.getBoundingClientRect();
                            const scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                            const targetY = rect.top + scrollTop - 80; // 상단 여백 80px
                            
                            window.scrollTo({
                                top: targetY,
                                behavior: 'smooth'
                            });
                            return true;
                        }
                        return false;
                    }
                    
                    // 페이지 로드 완료 후 실행
                    if (document.readyState === 'complete') {
                        scrollToViewer();
                    } else {
                        window.addEventListener('load', scrollToViewer);
                    }
                    
                    // 추가 재시도 (Streamlit 렌더링 지연 대응)
                    setTimeout(scrollToViewer, 500);
                    setTimeout(scrollToViewer, 1000);
                    setTimeout(scrollToViewer, 1500);
                })();
                </script>
                """,
                unsafe_allow_html=True,
            )
            st.session_state["recruit_admin_scroll_to_viewer"] = False
        
        # 닫기 버튼 추가
        col_title, col_close = st.columns([4, 1])
        with col_title:
            st.markdown(f"### 📄 공고 보기 - {rec_view.get('title')}")
        with col_close:
            if st.button("✕ 닫기", use_container_width=True, key="close_view"):
                # 닫기 전에 해당 채용공고 ID 저장
                view_rec_id = rec_view.get("id")
                if view_rec_id:
                    st.session_state["recruit_admin_scroll_to_rec_id"] = view_rec_id
                st.session_state["recruit_admin_view"] = None
                st.session_state["recruit_admin_list_expanded"] = True  # 목록 다시 펼치기
                st.rerun()
        
        path = Path(rec_view.get("file_path", ""))
        if path.exists():
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                try:
                    data = path.read_bytes()
                    b64 = base64.b64encode(data).decode("utf-8")
                    st.markdown(
                        f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="700px" style="border:none;border-radius:12px;box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);"></iframe>',
                        unsafe_allow_html=True,
                    )
                except Exception as e:
                    st.error(f"PDF 로드 실패: {e}")
            elif suffix == ".docx":
                try:
                    # DOCX 파일을 읽어서 HTML로 포맷팅
                    doc = Document(str(path))
                    html_content = '<div style="background: white; padding: 24px; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1); max-height: 700px; overflow-y: auto; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; line-height: 1.6; color: #1e293b;">'
                    
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if not text:
                            html_content += '<p style="margin: 8px 0;"></p>'
                            continue
                        
                        # 단락 스타일 확인
                        style_name = para.style.name if para.style else "Normal"
                        is_heading = style_name.startswith("Heading") or para.style.name.startswith("Title")
                        
                        if is_heading:
                            level = 1
                            if "Heading" in style_name:
                                try:
                                    level = int(style_name.replace("Heading ", ""))
                                except:
                                    level = 1
                            font_size = {1: "2em", 2: "1.5em", 3: "1.3em", 4: "1.1em"}.get(level, "1em")
                            font_weight = "bold"
                            margin = {1: "24px 0 16px", 2: "20px 0 12px", 3: "16px 0 10px"}.get(level, "12px 0 8px")
                            html_content += f'<h{min(level, 6)} style="font-size: {font_size}; font-weight: {font_weight}; margin: {margin}; color: #0f172a;">{text}</h{min(level, 6)}>'
                        else:
                            # 일반 단락
                            html_content += f'<p style="margin: 8px 0; font-size: 1em; color: #334155;">{text}</p>'
                    
                    # 리스트 처리
                    for para in doc.paragraphs:
                        if para.style.name.startswith("List"):
                            text = para.text.strip()
                            if text:
                                html_content += f'<li style="margin: 4px 0; padding-left: 8px;">{text}</li>'
                    
                    html_content += '</div>'
                    
                    st.markdown(html_content, unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"DOCX 로드 실패: {e}")
                    # 폴백: 텍스트로 표시
                    try:
                        doc = Document(str(path))
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        st.text_area("원문 텍스트", value=text, height=400, label_visibility="collapsed")
                    except:
                        st.error("파일을 읽을 수 없습니다.")
            else:
                try:
                    text = Path(path).read_text(encoding="utf-8", errors="ignore")
                    st.markdown(
                        f"""
                        <div style="background: white; padding: 24px; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1); max-height: 700px; overflow-y: auto; font-family: monospace; white-space: pre-wrap; line-height: 1.6; color: #1e293b;">
                        {text}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                except Exception as e:
                    st.error(f"파일 로드 실패: {e}")
        else:
            st.warning("파일을 찾을 수 없습니다.")

    # 업로드 폼 (expander 사용 - Streamlit 네이티브 방식)
    show_upload = st.session_state.get("recruit_admin_show_upload", False)
    
    if show_upload:
        st.markdown("---")
        # 스크롤 이동을 위한 앵커 및 스크립트
        st.markdown(
            """
            <div id="upload-form-anchor"></div>
            <script>
            (function() {
                function scrollToForm() {
                    const anchor = document.getElementById('upload-form-anchor');
                    if (anchor) {
                        anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                        return true;
                    }
                    return false;
                }
                // 즉시 시도
                if (!scrollToForm()) {
                    // DOM이 준비되지 않았으면 재시도
                    setTimeout(scrollToForm, 200);
                    setTimeout(scrollToForm, 500);
                }
            })();
            </script>
            """,
            unsafe_allow_html=True,
        )
        
        with st.expander("➕ 새 채용공고 등록", expanded=True):
            with st.form("upload_recruit", clear_on_submit=False):
                title = st.text_input("채용공고명 *", "")
                company = st.text_input("회사명", "미정")
                
                col_job1, col_job2 = st.columns(2)
                with col_job1:
                    job_family = st.text_input("직군", "")
                with col_job2:
                    role_category = st.text_input("직무", "")
                
                col_emp, col_exp = st.columns(2)
                with col_emp:
                    employment_type = st.selectbox("고용형태", ["정규", "계약", "인턴", "아르바이트", "기타"], index=0)
                with col_exp:
                    experience_level = st.selectbox("경력", ["무관", "신입", "경력"], index=0)
                
                location = st.text_input("근무지", "")
                
                col1, col2 = st.columns(2)
                with col1:
                    start_date = st.text_input("채용 시작일 (YYYY-MM-DD)", "")
                with col2:
                    end_date = st.text_input("채용 종료일 (YYYY-MM-DD)", "")
                
                status = st.selectbox("상태", ["OPEN", "CLOSED", "ARCHIVED"], index=0)
                file = st.file_uploader("채용공고 파일 업로드 *", type=["pdf", "docx", "txt", "md"])
                
                col_submit, col_cancel = st.columns([1, 1])
                with col_submit:
                    submitted = st.form_submit_button("✅ 등록", use_container_width=True, type="primary")
                with col_cancel:
                    cancel_clicked = st.form_submit_button("❌ 취소", use_container_width=True)

                if cancel_clicked:
                    st.session_state["recruit_admin_show_upload"] = False
                    st.session_state["recruit_admin_list_expanded"] = True  # 목록 다시 펼치기
                    st.rerun()

                if submitted:
                    if not title or not file:
                        st.error("채용공고명과 파일을 모두 입력해주세요.")
                    else:
                        # 필수 필드
                        data = {
                            "title": title.strip(),
                            "company": company.strip() if company else "미정",
                            "status": status,
                        }
                        # 선택적 필드 추가 (값이 있는 경우만)
                        if job_family and job_family.strip():
                            data["job_family"] = job_family.strip()
                        if role_category and role_category.strip():
                            data["role_category"] = role_category.strip()
                        if employment_type:
                            data["employment_type"] = employment_type
                        if experience_level:
                            data["experience_level"] = experience_level
                        if location and location.strip():
                            data["location"] = location.strip()
                        if start_date and start_date.strip():
                            data["start_date"] = start_date.strip()
                        if end_date and end_date.strip():
                            data["end_date"] = end_date.strip()
                        # posted_by는 int 또는 None (None이면 전송하지 않음)
                        member_id = st.session_state.get("member_id")
                        if member_id is not None:
                            data["posted_by"] = member_id
                        with st.spinner("등록 중..."):
                            result = upload_recruitment(data, file)
                        if result:
                            st.success("채용공고가 등록되었습니다.")
                            # 등록 영역 숨기기
                            st.session_state["recruit_admin_show_upload"] = False
                            # 목록 다시 펼치기
                            st.session_state["recruit_admin_list_expanded"] = True
                            # 상단으로 스크롤 이동 플래그 설정
                            st.session_state["recruit_admin_scroll_to_top"] = True
                            # 새로고침을 위한 플래그 설정
                            st.session_state["recruit_admin_last_upload"] = uuid.uuid4().hex
                            st.rerun()
